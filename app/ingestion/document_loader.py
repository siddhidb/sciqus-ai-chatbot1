# from pypdf import PdfReader
# from docx import Document
# import os


# def load_document(file_path: str) -> str:
#     ext = os.path.splitext(file_path)[1].lower()

#     if ext == ".pdf":
#         return _load_pdf(file_path)
#     if ext == ".docx":
#         return _load_docx(file_path)
#     if ext == ".txt":
#         return _load_txt(file_path)

#     raise ValueError(f"Unsupported document type: {ext}")


# def _load_pdf(path: str) -> str:
#     reader = PdfReader(path)
#     pages = []

#     for page in reader.pages:
#         text = page.extract_text()
#         if text:
#             pages.append(text)

#     return "\n".join(pages)


# def _load_docx(path: str) -> str:
#     doc = Document(path)
#     return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# def _load_txt(path: str) -> str:
#     with open(path, "r", encoding="utf-8", errors="ignore") as f:
#         return f.read()
# app/ingestion/document_loader.py

from pypdf import PdfReader
from docx import Document
import os


def load_document(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _load_pdf(file_path)
    elif ext == ".docx":
        return _load_docx(file_path)
    elif ext == ".txt":
        return _load_txt(file_path)

    raise ValueError(f"Unsupported document type: {ext}")


def _load_pdf(path: str) -> str:
    reader = PdfReader(path)
    pages = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)

    return "\n".join(pages)


def _load_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _load_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
